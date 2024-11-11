# -*- coding: utf-8 -*-
import streamlit as st
import os
import pandas as pd
import matplotlib.pyplot as plt
from resume_generation_gemini_pro import generate_gemini
from similarity_score_refined import similarity_main
from pdf2image import convert_from_path, convert_from_bytes
from docx import Document
import subprocess
import shutil
import io
from io import BytesIO
import tempfile
from PIL import Image, ImageDraw, ImageFont
import PyPDF2
from docx2pdf import convert
import pdfplumber
import docx
import numpy as np
import pypandoc
import streamlit.components.v1 as components

# Create temporary directories
temp_dir = tempfile.mkdtemp()

# Custom CSS for styling
st.markdown("""
    <style>
        .main {
            background-color: #f5f5f5;
            font-family: Arial, sans-serif;
        }
        h1, h2 {
            color: #4B7BE5;
            text-align: center;
        }
        .stContainer {
            # background-color: #000000;
            display: flex;
            justify-content: center;
            align-items: center;
            # max-width: 100%;
            height: 30%;
            width: 45%;
        }
        .logo-container {
            # background-color: black;
            display: flex;
            justify-content: center;
            align-items: center;
            padding: 10px;
            # max-width: 100%;
            
        }
        .logo-container img {
            max-width: 60%;
            height: 40%;
        }
        .stButton>button {
            # background-color: #4B7BE5;
            # color: white;
            # font-size: 18px;
            appearance: none;
            background-color: transparent;
            border: 0.125em solid #1A1A1A;
            border-radius: 0.9375em;
            box-sizing: border-box;
            color: #3B3B3B;
            cursor: pointer;
            display: inline-block;
            font-family: Roobert,-apple-system,BlinkMacSystemFont,"Segoe UI",Helvetica,Arial,sans-serif,"Apple Color Emoji","Segoe UI Emoji","Segoe UI Symbol";
            font-size: 16px;
            font-weight: 600;
            line-height: normal;
            margin: 0;
            min-height: 3.75em;
            min-width: 0;
            outline: none;
            padding: 1em 2.3em;
            text-align: center;
            text-decoration: none;
            transition: all 300ms cubic-bezier(.23, 1, 0.32, 1);
            user-select: none;
            -webkit-user-select: none;
            touch-action: manipulation;
            will-change: transform;
        }
        .stButton>button:hover {
            color: #fff;
            background-color: #1A1A1A;
            box-shadow: rgba(0, 0, 0, 0.25) 0 8px 15px;
            transform: translateY(-2px);
            border: none !important;            
            
        }
        /* From Uiverse.io by e-coders */ 
             
        # .stButton>btn:disabled {
        #  pointer-events: none;
        # }
        
        .stButton>:active, focus {
         box-shadow: none;
         transform: translateY(0);
         color: #fff;
         border: none !important;  
         outline: none;
        }
    </style>
""", unsafe_allow_html=True)

# Add ResumeMagic Logo
st.markdown('<div class="logo-container"></div>', unsafe_allow_html=True)
st.image("logo.jpeg", width=80)
st.markdown('</div>', unsafe_allow_html=True)


# Title and Description
st.title("Resume Tailoring with Google Generative AI")
st.markdown("### Upload your resume and job description to check similarity and generate a tailored resume.")


# Helper function to save uploaded files temporarily and return their paths
def save_uploaded_file(content):
    if hasattr(content, 'name'):  
        file_path = os.path.join("/tmp", content.name)
        with open(file_path, "wb") as f:
            f.write(content.read())
    else: 
        file_path = os.path.join("/tmp", "temp_upload")
        with open(file_path, "w") as f:
            f.write(str(content))
    return file_path
    
# def save_uploaded_file(uploaded_file):
#     file_path = os.path.join("/tmp", uploaded_file.name)
#     with open(file_path, "wb") as f:
#         f.write(uploaded_file.getbuffer())
#     return file_path


# Two columns for file uploaders
col1, col2 = st.columns(2)
with col1:
    uploaded_resume = st.file_uploader("Upload Current Resume (.docx or .pdf)", type=["docx", "pdf"], key="resume")
with col2:
    uploaded_job_description = st.file_uploader("Upload Job Description (.docx or .pdf)", type=["docx", "pdf"], key="job_description")

def get_score(resume_path, job_description_path):
    similarity_score = similarity_main(resume_path, job_description_path)
    if isinstance(similarity_score, str) and '%' in similarity_score:
        similarity_score = float(similarity_score.replace('%', ''))
        
        # Display messages based on score range
        if similarity_score < 50:
            st.markdown('<p style="color: red; font-weight: bold;">Low chance, skills gap identified!</p>', unsafe_allow_html=True)
            pie_colors = ['#FF4B4B', '#E5E5E5']  
        elif 50 <= similarity_score < 70:
            st.markdown('<p style="color: red; font-weight: bold;">Good chance but you can improve further!</p>', unsafe_allow_html=True)
            pie_colors = ['#FFC107', '#E5E5E5']  
        else:
            st.markdown('<p style="color: green; font-weight: bold;">Excellent! You can submit your CV.</p>', unsafe_allow_html=True)
            pie_colors = ['#4CAF50', '#E5E5E5']  
    
    return similarity_score, pie_colors

def display_score(similarity, colors):
    # Display Score as a Pie Chart
    st.markdown(f"### Resume - Job Match: {int(similarity_score)}%")
    
    # Pie chart to show similarity
    fig, ax = plt.subplots()
    # ax.pie([similarity_score, 100 - similarity_score], labels=['Match', 'Difference'], autopct='%1.1f%%', startangle=140, colors=['#4B7BE5', '#E5E5E5'])
    ax.pie([similarity_score, 100 - similarity_score], labels=['Match', 'Difference'], autopct='%1.1f%%', startangle=140, colors=pie_colors)
        
    ax.axis('equal')  
    st.pyplot(fig)

    
def display_docx_content(file):
    doc = docx.Document(file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)
    

# Function to save a file from BytesIO to a temporary file
def save_bytes_to_tempfile(bytes_data, suffix):
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
        temp_file.write(bytes_data)
        return temp_file.name
        
def save_bytes_as_pdf(docx_bytes, output_path='output.pdf'):
    # Create a temporary directory
    with tempfile.TemporaryDirectory() as tmp_dir:
        # Write the DOCX bytes to a temporary file
        temp_file = os.path.join(tmp_dir, 'temp.docx')
        with open(temp_file, 'wb') as f:
            f.write(docx_bytes)
        
        # Convert the temporary DOCX to PDF
        pdf_path = os.path.join(tmp_dir, 'output.pdf')
        convert(temp_file, pdf_path)
        
        # Copy the PDF to the desired output location
        with open(output_path, 'wb') as f:
            with open(pdf_path, 'rb') as src_f:
                f.write(src_f.read())

    # Clean up the temporary directory
    os.remove(output_path)


def display_content_with_page_numbers(content, words_per_page=290):
    # Split content into words
    words = content.split()
    total_pages = (len(words) // words_per_page) + (1 if len(words) % words_per_page != 0 else 0)

    # Display content with page numbers
    for i in range(total_pages):
        start_index = i * words_per_page
        end_index = start_index + words_per_page
        page_content = ' '.join(words[start_index:end_index])
        
        # st.markdown(f"#### Page {i + 1}")
        # st.write(page_content)
    st.markdown(f"#### Page {total_pages}")
    
def save_docx_as_pdf(input_path, output_path='output.pdf'):
    if input_path.lower().endswith('.docx'):
        try:
            # Convert .docx to .pdf using LibreOffice
            subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', input_path, '--outdir', os.path.dirname(output_path)], check=True)
            if not os.path.exists(output_path):
                raise FileNotFoundError("Conversion failed; output PDF not found.")
        except (FileNotFoundError, subprocess.CalledProcessError):
            st.error("Failed to convert DOCX to PDF. Please check LibreOffice installation.")
    elif input_path.lower().endswith('.pdf'):
        shutil.copy(input_path, output_path)
    else:
        raise ValueError("Unsupported file format. Please upload a .docx or .pdf file.")


def display_pdf_page(pdf_path):
    try:
        # Open PDF file
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
        
        # Extract text from the first page
        page = reader.pages[0]
        x_object = page.extract_text()
    
        # Convert text to image (using PIL)
        img = Image.new('RGB', (800, 1000))
        draw = ImageDraw.Draw(img)
        font = ImageFont.truetype("arial.ttf", 20)
    
        # Draw text on the image
        draw.text((10, 10), x_object[:500], fill=(255, 255, 255), font=font)
    
        # Display the image
        display(img)
    except Exception as e:
        st.error(f"Failed to display image: {str(e)}")

# def display_pdf_pages_as_images(pdf_path):
#     try:
#         with pdfplumber.open(pdf_path) as pdf:
#             for i, page in enumerate(pdf.pages):
#                 st.markdown(f"#### Page {i + 1}")
#                 # Convert the page to an image
#                 image = page.to_image()
#                 # Render the image using Streamlit
#                 # st.image(image.original, use_column_width=True)
#                 st.image(image.original, use_container_width=False)
#     except Exception as e:
#         st.error(f"Failed to display PDF as image: {str(e)}")

def display_pdf_pages_as_images(pdf_path):
    try:
        with pdfplumber.open(pdf_path) as pdf:
            num_pages = len(pdf.pages)
            # Create a container with columns for each page
            columns = st.columns(num_pages)

            for i, page in enumerate(pdf.pages):
                # Convert the page to an image
                image = page.to_image()
                # Display each page image in its respective column
                with columns[i]:
                    st.markdown(f"#### Page {i + 1}")
                    st.image(image.original, use_container_width=True)
                    
    except Exception as e:
        st.error(f"Failed to display PDF as image: {str(e)}")

    
def display_doc_as_image2(pdf_path):   
    iframe_code = f"""
    <iframe src="{pdf_path}" width="100%" height="600px"></iframe>
    """
    st.markdown(iframe_code, unsafe_allow_html=True)
    
    
 
# Process if files are uploaded
if uploaded_resume and uploaded_job_description:
    # Save files
    resume_path = save_uploaded_file(uploaded_resume)
    job_description_path = save_uploaded_file(uploaded_job_description)

    # Similarity Score Section
    st.markdown("---")
    # st.subheader("Check Job Match")

    if st.button("Resume-JD Matching"):
        with st.spinner("Computing Match"):
            similarity_score, pie_colors = get_score(resume_path, job_description_path)
            display_score(similarity_score, pie_colors)
 
        #Autoscroll
        components.html("""
                <script>
                    window.onload = function() {
                        window.scrollTo(0, document.body.scrollHeight);
                    };
                </script>
            """)

    # Generate Tailored Resume Section
    st.markdown("---")
    # st.subheader("Tailor Resume")
    
    if st.button("Tailor Resume"):
        with st.spinner("Generating resume..."):
            generated_resume, new_resume_path = generate_gemini(resume_path, job_description_path)
            # resume_path = save_uploaded_file(generated_resume)
            # st.markdown("Generated Tailored Resume:")
            # st.write(generated_resume)

            #Autoscroll
            components.html("""
                <script>
                    window.onload = function() {
                        window.scrollTo(0, document.body.scrollHeight);
                    };
                </script>
            """)

            
            # with st.spinner("Computing Match"):
            #     similarity_score, pie_colors = get_score(resume_path, job_description_path)
            #     display_score(similarity_score, pie_colors)
                
            if generated_resume is not None:
                st.markdown("---")
                st.title("Uploaded Resumes")
                doc = Document()
                doc.add_paragraph(generated_resume)  
                
                # Save the generated document as a .docx file in memory
                resume_bytes = BytesIO()
                doc.save(resume_bytes)
                resume_bytes.seek(0)
                
                # Save the .docx to a temporary file
                gen_docx_path = save_bytes_to_tempfile(resume_bytes.getvalue(), 'docx')
        
                # Convert the generated .docx to a .pdf
                gen_pdf_path = save_uploaded_file(gen_docx_path)
                # st.write(display_docx_content(gen_pdf_path))

                
                # st.markdown("### Uploaded Resume")
                save_docx_as_pdf(resume_path, '/tmp/uploaded_resume.pdf')
                display_pdf_pages_as_images(resume_path)
               
                st.success(f"Download tailored resume")
                col1, col2 = st.columns(2)
                with col1:
                    st.download_button(
                        label="Generated Resume (PDF)",
                        data=open(gen_pdf_path, 'rb').read(),
                        file_name="tailored_resume.pdf",
                        mime="application/pdf"
                    )
                with col2:
                    st.download_button(
                        label="Generated Resume (Word)",
                        data=resume_bytes,
                        file_name="tailored_resume.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )                            


                # Display uploaded and generated resumes side-by-side
                # col1, col2 = st.columns(2)
                # with col1:
                #     st.markdown("### Uploaded Resume")
                #     save_docx_as_pdf(resume_path, '/tmp/uploaded_resume.pdf')
                #     display_pdf_pages_as_images(resume_path)
                # with col2:
                #     st.markdown("### Tailored Resume")
                #     # display_pdf_pages_as_images(gen_pdf_path)
                #     display_content_with_page_numbers(generated_resume, 290)
                #     st.write(generated_resume)
                #     # display_content_with_page_numbers(generated_resume, 290)
                    
                #     st.success(f"Download tailored resume")
                #     col1, col2 = st.columns(2)
                #     with col1:
                #         st.download_button(
                #             label="Generated Resume (PDF)",
                #             data=open(gen_pdf_path, 'rb').read(),
                #             file_name="tailored_resume.pdf",
                #             mime="application/pdf"
                #             )
                #     with col2:
                #         st.download_button(
                #             label="Generated Resume (Word)",
                #             data=resume_bytes,
                #             file_name="tailored_resume.docx",
                #             mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                #         )                            
                
                
else:
    st.warning("Please upload both the resume and job description files.")

# -*- coding: utf-8 -*-
"""Resume_generation_Gemini_pro.ipynb

Automatically generated by Colab.

Original file is located at
    https://colab.research.google.com/drive/16z793IRwVmvKYCaOLGZFDYj-XOj8zEJL
"""

# from google.colab import drive,userdata

# drive.mount('/content/drive')

# !pip install streamlit -qq

# !pip install PyPDF2 -qq

# !pip install langchain_community -qq

# !pip install langchain_google_genai -qq

# !pip install python-docx -qq

# !pip install docx2txt -qq

# !pip install faiss-gpu -qq

# !pip install google-generativeai -qq

# !pip install --upgrade google-generativeai -qq

# !pip install python-docx

import docx2txt
import PyPDF2
import subprocess
def extract_text(file_path):
    if file_path.endswith(".docx"):
        # Extract text from DOCX file
        return docx2txt.process(file_path)

    elif file_path.endswith(".pdf"):
        # Extract text from PDF file
        text = ""
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page_num in range(len(reader.pages)):
                text += reader.pages[page_num].extract_text()
        return text

    else:
        raise ValueError("Unsupported file type")

import os
# import streamlit as st
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores.faiss import FAISS
# from google.colab import drive
from docx import Document
import google.generativeai as genai
from datetime import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from dotenv import load_dotenv
load_dotenv()

api_key_google = os.environ.get('GOOGLE_GEMINI_KEY')
genai.configure(api_key=api_key_google)


model = genai.GenerativeModel('gemini-pro')

def save_resume_to_docx(tailored_resume, file_path):
    doc = Document()
    doc.add_heading('Tailored Resume', level=1)
    doc.add_paragraph(tailored_resume)
    doc.save(file_path)

def save_resume_to_pdf(docx_file_path, file_path):
    subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', file_path, docx_file_path])


# Function to read text from a .docx file
def read_docx(file_path):
    doc = Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

# def generate_resume_text(resume_text):
#     prompt = f"""
# Given the following resume content:

# [Resume Start]
# {resume_text}
# [Resume End]

# Format this resume content with appropriate section titles. Only use the information provided and avoid placeholders like "[Your Name]". Ensure it retains the structure and details exactly as shown.
# """
#     try:
#         response = model.generate_content(prompt)
#         print(response)
#         # Accessing the generated text content
#         return response.candidates[0].content.parts[0].text
#     except Exception as e:
#         print("Error in generating resume text:", e)
#         return None

def tailor_resume(resume_text, job_description):
    # Use the generate_resume_text function to get the formatted resume content
    # formatted_resume = generate_resume_text(resume_text)
    # print("formatted resume:",resume_text)
        prompt = f"""
Below is the candidate's original  resume content:
[Resume Start]
{resume_text}
[Resume End]
Using the candidate's resume above and the job description below, create a tailored resume.
[Job Description Start]
{job_description}
[Job Description End]
Please generate a resume that:
1. Uses real data from the candidate's resume, including name, and education.
2. Avoids placeholders like "[Your Name]" and includes actual details. This is important.
3. In the experience section, emphasizes professional experiences and skills that are directly relevant to the job description.
4. Keeps only a maximum of the top three accomplishments/ responsibilities for each job position held so as to make the candidate standout in the new job role
5. Removes special characters from the section titles.
6. Only includes publications if the job description is research based.
7. Summarizes the skills and technical skills section into a brief profile.
8. Does not include courses, certification, references, skills and a technical skills sections.
9. Only includes true information about the candidate.It is very important that no fake details are added.Do not add anything that candidate has not worked on or has expereince with.
10.If profile or summary section is being is being included, give that section priority after candiadte's information.
11.Try to keep the content to one page.
12.If education is the strongest asset of the candidate, give it a priority over experience and vice-versa.
13.Provide the text in markdown format that clearly identifies the headings and subheadings.
"""


        try:
            response = model.generate_content(prompt)
            print(response.candidates[0].content.parts[0].text)
            return response.candidates[0].content.parts[0].text
        except Exception as e:
            print("Error in tailoring resume:", e)
            return None

def add_bold_and_normal_text(paragraph, text):
    """Adds text to the paragraph, handling bold formatting."""
    while "**" in text:
        before, bold_part, after = text.partition("**")
        if before:
            paragraph.add_run(before)
        if bold_part == "**":
            bold_text, _, text = after.partition("**")
            paragraph.add_run(bold_text).bold = True
        else:
            text = after
    if text:
        paragraph.add_run(text)
        
def convert_resume_to_word(markdown_text,output_file):
    # Create a new Word document
    doc = Document()

    # Split the text into lines for processing
    lines = markdown_text.splitlines()

    for line in lines:
                  if line.startswith("# "):  # Top-level heading (Highest level)
                    paragraph = doc.add_heading(line[2:].strip(), level=0)  # Level 0 is the highest heading in Word
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  
                  elif line.startswith("## "):  # Main heading (Level 1)
                    paragraph = doc.add_heading(line[3:].strip(), level=1)
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                  elif line.startswith("### "):  # Subheading (Level 2)
                    paragraph = doc.add_heading(line[4:].strip(), level=2)
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                  elif line.startswith("- "):  # Bullet points
                    paragraph = doc.add_paragraph(style="List Bullet")
                    add_bold_and_normal_text(paragraph, line[2:].strip())
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                  elif line.startswith("* "):  # Sub-bullet points or normal list items
                    paragraph = doc.add_paragraph(style="List Bullet 2")
                    add_bold_and_normal_text(paragraph, line[2:].strip())
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                  elif line.strip():  # Normal text (ignores blank lines)
                    paragraph = doc.add_paragraph()
                    add_bold_and_normal_text(paragraph, line.strip())
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Save the Word document

    doc.save(output_file)
    print(f"Markdown converted and saved as {output_file}")

#Entry function for the model
def generate_gemini(current_resume,job_description , download_path):
    # st.header('Resume Tailoring')
    # Load the resume and job description from Google Drive
    resume_text = extract_text(current_resume)
    job_description = extract_text(job_description)

    # Tailor resume based on job description
    tailored_resume = tailor_resume(resume_text, job_description)
    # st.write("**Tailored Resume:**")
    # st.write(tailored_resume)
    print(tailored_resume)

        # Save the tailored resume to a .docx file
    if tailored_resume:
        time = datetime.now().strftime('%Y%m%d_%H%M%S')
        # Generate the docx file
        file_name = f"Tailored_Resume_{time}.docx"
        file_path = os.path.join(download_path , file_name)
        convert_resume_to_word(tailored_resume,file_path)
        
        # Generate the pdf file
        file_name = f"Tailored_Resume_{time}.pdf"
        save_resume_to_pdf(file_path, download_path)

        # st.success(f"Download tailored resume")
        # st.success(f"Tailored resume saved to {file_path}")        

    return tailored_resume, time

def returnFile(file_name , filetype):
    resume_path = f"Tailored_Resume_{file_name}.{filetype}"
    return resume_path
